from __future__ import annotations

from pathlib import Path

from docx import Document

from .models import SearchResultPackage
from .utils import dump_json, ensure_dir, journal_scope_label, normalize_requested_journals


def _data_source_label(package: SearchResultPackage) -> str:
    if package.config.enable_semantic_scholar:
        return "DBLP + Crossref + Semantic Scholar（仅 enrich）"
    return "DBLP + Crossref"


def render_markdown(package: SearchResultPackage) -> str:
    config = package.config
    stats = package.stats
    journals = normalize_requested_journals(config.normalized_journals())
    scope_label = journal_scope_label(journals)
    data_sources = _data_source_label(package)

    if config.days_back is not None:
        time_window = f"最近 {config.days_back} 天"
    elif config.years_back is not None:
        time_window = f"最近 {config.years_back} 年"
    else:
        time_window = "未限定"

    lines: list[str] = [
        "# 机器人文献检索总结",
        "",
        "## 检索任务说明",
        f"- 检索时间：{package.generated_at}",
        f"- 来源范围：{scope_label}",
        f"- 数据源：{data_sources}",
        f"- 关键词：{', '.join(config.keywords)}",
        f"- 时间范围：{time_window}",
        f"- 返回论文数量：{len(package.papers)}",
        "",
        "## 总体观察",
        package.overview,
        "",
        "## 单篇论文整理",
        "",
    ]

    if not package.papers:
        lines.append("本次没有符合条件的论文。")
    else:
        for idx, paper in enumerate(package.papers, start=1):
            lines.extend(
                [
                    f"### {idx}. {paper.title}",
                    f"- 作者：{', '.join(paper.authors) if paper.authors else '未知'}",
                    f"- 日期 / 年份：{paper.publication_date or paper.publication_year or '未知'}",
                    f"- DOI：{paper.doi or '无'}",
                    f"- 链接：{paper.url or paper.dblp_url or '无'}",
                    f"- Venue：{paper.journal or paper.raw_venue or '未知'}",
                    f"- 命中关键词：{', '.join(paper.matched_keywords) if paper.matched_keywords else '无'}",
                    f"- 主题标签：{', '.join(paper.topic_tags) if paper.topic_tags else '无'}",
                    f"- 摘要来源：{paper.abstract_source or '无'}",
                    f"- 原始摘要：{paper.abstract or '未获取到公开摘要'}",
                    f"- 摘要中文翻译：{paper.chinese_summary or '未获取到公开摘要，暂无可用的摘要中文翻译。'}",
                    "",
                ]
            )

    lines.extend(
        [
            "## 关键词分组观察",
            package.keyword_observations,
            "",
            "## 推荐优先阅读列表",
            "",
        ]
    )

    if not package.recommendations:
        lines.append("暂无推荐论文。")
    else:
        for idx, item in enumerate(package.recommendations, start=1):
            lines.extend(
                [
                    f"{idx}. {item.get('title', '未知标题')}",
                    f"   - DOI：{item.get('doi') or '无'}",
                    f"   - 年份：{item.get('publication_year') or '未知'}",
                    f"   - 排序分：{item.get('rank_score')}",
                ]
            )

    lines.extend(
        [
            "",
            "## 统计信息",
            f"- DBLP 候选数：{stats.dblp_candidates}",
            f"- Crossref 补全数：{stats.crossref_enriched}",
            f"- Semantic Scholar 补全数：{stats.semantic_scholar_enriched}",
            f"- 最终论文数：{stats.final_papers}",
            f"- 有摘要论文数：{stats.papers_with_abstract}",
            f"- 有 DOI 论文数：{stats.papers_with_doi}",
            f"- 有开放获取 PDF 链接的论文数：{stats.papers_with_pdf}",
        ]
    )

    return "\n".join(lines).strip() + "\n"


def save_markdown_report(package: SearchResultPackage, path: str | Path) -> Path:
    out_path = Path(path)
    out_path.write_text(render_markdown(package), encoding="utf-8")
    return out_path


def save_docx_report(package: SearchResultPackage, path: str | Path) -> Path:
    doc = Document()
    config = package.config
    stats = package.stats
    journals = normalize_requested_journals(config.normalized_journals())
    scope_label = journal_scope_label(journals)
    data_sources = _data_source_label(package)

    doc.add_heading("机器人文献检索总结", level=0)

    doc.add_heading("检索任务说明", level=1)
    doc.add_paragraph(f"检索时间：{package.generated_at}")
    doc.add_paragraph(f"来源范围：{scope_label}")
    doc.add_paragraph(f"数据源：{data_sources}")
    doc.add_paragraph(f"关键词：{', '.join(config.keywords)}")
    if config.days_back is not None:
        doc.add_paragraph(f"时间范围：最近 {config.days_back} 天")
    elif config.years_back is not None:
        doc.add_paragraph(f"时间范围：最近 {config.years_back} 年")
    else:
        doc.add_paragraph("时间范围：未限定")
    doc.add_paragraph(f"返回论文数量：{len(package.papers)}")

    doc.add_heading("总体观察", level=1)
    for paragraph in package.overview.split("\n\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())

    doc.add_heading("单篇论文整理", level=1)
    if not package.papers:
        doc.add_paragraph("本次没有符合条件的论文。")
    else:
        for idx, paper in enumerate(package.papers, start=1):
            doc.add_heading(f"{idx}. {paper.title}", level=2)
            doc.add_paragraph(f"作者：{', '.join(paper.authors) if paper.authors else '未知'}")
            doc.add_paragraph(f"日期 / 年份：{paper.publication_date or paper.publication_year or '未知'}")
            doc.add_paragraph(f"DOI：{paper.doi or '无'}")
            doc.add_paragraph(f"链接：{paper.url or paper.dblp_url or '无'}")
            doc.add_paragraph(f"Venue：{paper.journal or paper.raw_venue or '未知'}")
            doc.add_paragraph(f"命中关键词：{', '.join(paper.matched_keywords) if paper.matched_keywords else '无'}")
            doc.add_paragraph(f"主题标签：{', '.join(paper.topic_tags) if paper.topic_tags else '无'}")
            doc.add_paragraph(f"摘要来源：{paper.abstract_source or '无'}")
            doc.add_paragraph(f"原始摘要：{paper.abstract or '未获取到公开摘要'}")
            doc.add_paragraph(
                f"摘要中文翻译：{paper.chinese_summary or '未获取到公开摘要，暂无可用的摘要中文翻译。'}"
            )

    doc.add_heading("关键词分组观察", level=1)
    for line in package.keyword_observations.splitlines():
        doc.add_paragraph(line)

    doc.add_heading("推荐优先阅读列表", level=1)
    if not package.recommendations:
        doc.add_paragraph("暂无推荐论文。")
    else:
        for idx, item in enumerate(package.recommendations, start=1):
            doc.add_paragraph(
                f"{idx}. {item.get('title', '未知标题')} | "
                f"年份：{item.get('publication_year') or '未知'} | "
                f"DOI：{item.get('doi') or '无'}"
            )

    doc.add_heading("统计信息", level=1)
    doc.add_paragraph(f"DBLP 候选数：{stats.dblp_candidates}")
    doc.add_paragraph(f"Crossref 补全数：{stats.crossref_enriched}")
    doc.add_paragraph(f"Semantic Scholar 补全数：{stats.semantic_scholar_enriched}")
    doc.add_paragraph(f"最终论文数：{stats.final_papers}")
    doc.add_paragraph(f"有摘要论文数：{stats.papers_with_abstract}")
    doc.add_paragraph(f"有 DOI 论文数：{stats.papers_with_doi}")
    doc.add_paragraph(f"有开放获取 PDF 链接的论文数：{stats.papers_with_pdf}")

    out_path = Path(path)
    doc.save(out_path)
    return out_path


def write_outputs(
    package: SearchResultPackage,
    run_dir: str | Path,
    save_json_flag: bool,
    save_markdown_flag: bool,
    save_docx_flag: bool,
) -> dict[str, str]:
    output_dir = ensure_dir(run_dir)
    generated_files: dict[str, str] = {}

    if save_markdown_flag:
        md_path = output_dir / "summary.md"
        save_markdown_report(package, md_path)
        generated_files["summary_md"] = str(md_path)

    if save_docx_flag:
        docx_path = output_dir / "summary.docx"
        save_docx_report(package, docx_path)
        generated_files["summary_docx"] = str(docx_path)

    if save_json_flag:
        json_path = output_dir / "results.json"
        generated_files["results_json"] = str(json_path)

    package.stats.generated_files = generated_files

    if save_json_flag:
        dump_json(output_dir / "results.json", package.to_dict())

    return generated_files
